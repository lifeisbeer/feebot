from docx import Document
from os import walk

# flags to control the operation of the program
# ! set this to true to create the mark schemes
create_markschemes = True 
# ! set this to true to create the feedback
create_feedback = True

# read correct feedback from file
with open('feedback_correct.txt') as f:
    lines = f.readlines()
correct_feedback = []
for l in lines:
    correct_feedback += [l[:-1]]
#print(correct_feedback)

# read wrong feedback from file
with open('feedback_wrong.txt') as f:
    lines = f.readlines()
wrong_feedback = []
for l in lines:
    wrong_feedback += [l[:-1]]
#print(wrong_feedback)

# read missing feedback from file
with open('feedback_missing.txt') as f:
    lines = f.readlines()
missing_feedback = []
for l in lines:
    missing_feedback += [l[:-1]]
#print(missing_feedback)

# create a list with all the filenames of pdfs in the current directory
files = []
filenames = next(walk("./"), (None, None, []))[2] # This gives you all the files in the current directory
for f in filenames:
    f = f.split('.')
    if len(f) == 2 and f[1] == 'pdf':
        files += [f[0]]
#print(files)

#files = ["0test", "1test"] # !!! for test remove later !!!

if create_markschemes:
    mark_scheme = Document('mark_scheme.docx')
    for f in files:
        if not f + '.docx' in filenames:
            mark_scheme.save(f + '.docx')
            #print('Created mark scheme for ' + f)

if create_feedback:
    feedback = Document()
    # for each file create feedback
    for f in files:
        #feedback = Document()
        feedback.add_heading(f, level=1)
        mark_sheet = Document(f+'.docx')
        paragraphs = mark_sheet.paragraphs

        total = 0
        score = 0
        num = 0 # keep track of the questions
        try:
            for p in paragraphs:
                text = p.text
                if text[0] == 'Q': # for marks
                    #print(text)
                    # add question and mark in feedback
                    paragraph = feedback.add_paragraph(text+' - ')
                    # increase score
                    split_text = text.split(':')
                    mark_text = split_text[1].split('/')
                    score += int(mark_text[0])
                    total += int(mark_text[1])      
                else: # for feedback
                    #print(p.text)
                    runs = p.runs
                    for r in runs:
                        if r.font.highlight_color == 7: # add correct feedback and break
                            #print(correct_feedback[num])
                            paragraph.add_run(correct_feedback[num])
                            num += 1  
                            break
                        elif r.font.highlight_color == 6: # add incorrect feedback and break
                            #print(wrong_feedback[num])
                            paragraph.add_run(wrong_feedback[num])
                            num += 1  
                            break  
                        elif r.font.highlight_color == 4: # add feedback directly from mark sheet
                            #print(r.text)
                            paragraph.add_run(r.text)              
                        elif r == runs[len(runs)-1]: # add missing feedback
                            #print(missing_feedback[num])
                            paragraph.add_run(missing_feedback[num])
                            num += 1 

            # add total score in feedback
            feedback.add_paragraph("Total: " + str(score) + '/' + str(total))
            #feedback.save(f+'_feedback.docx')
        except:
            print('Error in file ' + f)
    feedback.save('full_feedback.docx')